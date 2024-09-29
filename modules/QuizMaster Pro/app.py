import os
import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from io import BytesIO
from fpdf import FPDF  
from docx import Document  
from dotenv import load_dotenv
import re

# Load environment variables
load_dotenv()
google_api_key = os.getenv("GOOGLE_API_KEY")

# CheatSheetGenerator class to handle quiz generation
class QuizGenerator:
    def __init__(self):
        # Initialize the LLM with the Google Gemini API
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            temperature=0.2,
            max_tokens=None,
            api_key=google_api_key
        )

    def generate_quiz(self, topic, question_type):
        # Create a detailed prompt to generate quiz questions
        prompt = f"""
        Create a set of quiz questions on '{topic}' focusing on '{question_type}' type questions.
        - For **Multiple Choice (MCQ)**: Provide a question and four options with one correct answer.
        - For **True/False**: Provide a question and indicate whether the statement is true or false.
        - For **Fill in the Blanks**: Provide a statement with a blank space for users to fill in.
        
        Ensure the questions are concise and relevant to the topic. Clearly mention the correct answer for evaluation purposes.
        """

        messages = [("system", "You are a helpful assistant."), ("human", prompt)]
        response = self.llm.invoke(messages).content
        return response

# DocumentHandler class to handle export functionality
class DocumentHandler:
    def export_to_docx(self, quiz_text, filename):
        # Create a Word document with the generated quiz text
        doc = Document()
        doc.add_heading('Quiz', 0)
        doc.add_paragraph(quiz_text)

        # Save the document to a BytesIO object for downloading
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def export_to_pdf(self, quiz_text, filename):
        # Create a PDF document with the quiz text
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, quiz_text)

        # Save the PDF content to a BytesIO object for downloading
        buffer = BytesIO()
        pdf_output = pdf.output(dest='S').encode('latin1')  # Get the content as a byte string
        buffer.write(pdf_output)
        buffer.seek(0)  # Move the cursor to the start of the buffer
        return buffer

# Function to extract questions and answers from generated content
def extract_questions_and_answers(quiz_text):
    questions = []
    answers = []
    
    # Regex to find MCQs, True/False, and Fill in the Blanks
    question_regex = r"Q\d+:(.*?)\n"
    answer_regex = r"Answer:\s*(.*?)\n"
    
    # Find all questions
    question_matches = re.findall(question_regex, quiz_text, re.DOTALL)
    answer_matches = re.findall(answer_regex, quiz_text)

    # Store questions and answers
    for q, a in zip(question_matches, answer_matches):
        questions.append(q.strip())
        answers.append(a.strip())

    return questions, answers

# Streamlit application layout
def main():
    st.set_page_config(page_title="QuizMaster Pro", layout="wide")  # Set page title and layout

    st.title("🎓 QuizMaster Pro")
    st.sidebar.header("Generate and Take Your Quiz")  # Sidebar header

    st.markdown(""" 
        **QuizMaster Pro** helps you create practice quizzes instantly. 
        Just enter a topic or upload a document, select your preferred question type, and start practicing. 
        Ideal for preparing for exams or testing your knowledge!
    """)

    # Input field for topic in the sidebar
    topic = st.sidebar.text_input("Enter Topic:", placeholder="e.g., Machine Learning")

    # Option to upload a document in the sidebar
    uploaded_file = st.sidebar.file_uploader("Or upload a document (PDF/DOCX):", type=["pdf", "docx"])

    # Dropdown for selecting question type
    question_type = st.sidebar.selectbox(
        "Select Question Type:",
        ("Multiple Choice (MCQ)", "True/False", "Fill in the Blanks")
    )

    # Button to generate quiz in the sidebar
    if st.sidebar.button("Generate Quiz"):
        if topic or uploaded_file:
            with st.spinner("Generating quiz..."):
                generator = QuizGenerator()

                # If document is uploaded, process the document to extract content
                if uploaded_file:
                    if uploaded_file.type == "application/pdf":
                        from PyPDF2 import PdfReader
                        pdf_reader = PdfReader(uploaded_file)
                        topic = "".join([page.extract_text() for page in pdf_reader.pages])
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        doc = Document(uploaded_file)
                        topic = "\n".join([para.text for para in doc.paragraphs])

                # Generate quiz based on the topic or document content
                quiz = generator.generate_quiz(topic, question_type)

            # Store the quiz in session state to prevent it from disappearing
            st.session_state.quiz = quiz

    # Display the quiz if it exists and allow user to take the quiz
    if 'quiz' in st.session_state:
        st.subheader("Generated Quiz")
        st.write(st.session_state.quiz)

        # Extract questions and answers
        questions, answers = extract_questions_and_answers(st.session_state.quiz)

        # List to store user responses
        user_answers = []

        # Display each question and input options for answers
        for i, question in enumerate(questions):
            st.write(f"Q{i+1}: {question}")
            if "True/False" in question_type:
                user_response = st.radio(f"Answer for Q{i+1}", ("True", "False"))
            elif "Multiple Choice (MCQ)" in question_type:
                # Display each option on a separate line
                options = ["Option A", "Option B", "Option C", "Option D"]
                user_response = st.radio(f"Answer for Q{i+1}", options)
            else:  # Fill in the blanks
                user_response = st.text_input(f"Answer for Q{i+1}")

            user_answers.append(user_response)

        # Button to submit answers and evaluate the quiz
        if st.button("Submit Quiz"):
            # Calculate the score
            score = 0
            for user_answer, correct_answer in zip(user_answers, answers):
                if user_answer.lower() == correct_answer.lower():
                    score += 1

            # Display the score
            st.success(f"Your Score: {score}/{len(questions)}")

    else:
        st.warning("Please enter a topic or upload a document to generate the quiz.")

if __name__ == "__main__":
    main()
